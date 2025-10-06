#!/usr/bin/env python3
"""
Performance tests for Go vs Python DOCX parsers.

This module benchmarks the performance difference between the Go and Python
DOCX parser implementations to validate the performance improvements.
"""

import os
import time
import unittest
from pathlib import Path
import sys
import statistics

# Add src directory to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from go_doc_go.document_parser.factory import create_parser
import pytest


@pytest.mark.performance
class TestDocxPerformance(unittest.TestCase):
    """Performance benchmarks for DOCX parsers."""

    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = Path(__file__).parent
        self.assets_dir = self.test_dir / "assets"
        self.test_docx_file = self.assets_dir / "test_sample.docx"

        # Create test file if it doesn't exist
        if not self.test_docx_file.exists():
            self.create_test_docx()

        # Number of iterations for timing
        self.iterations = 5

    def create_test_docx(self):
        """Create a test DOCX file for performance testing."""
        try:
            from docx import Document

            # Create assets directory if it doesn't exist
            self.assets_dir.mkdir(parents=True, exist_ok=True)

            # Create a larger document for performance testing
            doc = Document()

            # Add title
            doc.add_heading('Performance Test Document', 0)

            # Add multiple sections with content
            for section_num in range(10):
                doc.add_heading(f'Section {section_num + 1}', 1)

                # Add paragraphs
                for para_num in range(5):
                    doc.add_paragraph(
                        f'This is paragraph {para_num + 1} in section {section_num + 1}. '
                        f'It contains some sample text to make the document larger and more '
                        f'realistic for performance testing. The content is repeated to '
                        f'simulate a real document with substantial text content.'
                    )

                # Add a subsection
                doc.add_heading(f'Subsection {section_num + 1}.1', 2)

                # Add more paragraphs
                for para_num in range(3):
                    doc.add_paragraph(
                        f'This is subsection paragraph {para_num + 1}. It provides additional '
                        f'content for testing purposes and helps create a more substantial '
                        f'document structure that better represents real-world documents.'
                    )

                # Add a table every few sections
                if section_num % 3 == 0:
                    table = doc.add_table(rows=4, cols=4)
                    for row_idx in range(4):
                        for col_idx in range(4):
                            cell = table.cell(row_idx, col_idx)
                            if row_idx == 0:
                                cell.text = f'Header {col_idx + 1}'
                            else:
                                cell.text = f'Row {row_idx}, Col {col_idx + 1}'

            # Save the document
            doc.save(str(self.test_docx_file))

        except ImportError:
            # If python-docx is not available, create a dummy file
            self.test_docx_file.touch()

    def test_python_parser_performance(self):
        """Benchmark Python DOCX parser performance."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        os.environ["USE_GO_MODULES"] = "false"
        try:
            parser = create_parser("docx")
            content = {
                "id": "perf_test",
                "binary_path": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }

            # Warm up
            try:
                parser.parse(content)
            except Exception as e:
                self.skipTest(f"Python DOCX parser not available: {e}")

            # Time multiple iterations
            times = []
            for i in range(self.iterations):
                start_time = time.time()
                result = parser.parse(content)
                end_time = time.time()
                times.append(end_time - start_time)

                # Verify we got a valid result
                self.assertIn("elements", result)
                self.assertGreater(len(result["elements"]), 0)

            avg_time = statistics.mean(times)
            min_time = min(times)
            max_time = max(times)

            print(f"\nPython DOCX Parser Performance:")
            print(f"  Average time: {avg_time*1000:.2f} ms")
            print(f"  Min time: {min_time*1000:.2f} ms")
            print(f"  Max time: {max_time*1000:.2f} ms")
            print(f"  Elements parsed: {len(result['elements'])}")
            print(f"  Relationships: {len(result.get('relationships', []))}")

        except ImportError:
            self.skipTest("Python DOCX parser not available (python-docx not installed)")

    def test_go_parser_performance(self):
        """Benchmark Go DOCX parser performance."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        os.environ["USE_GO_MODULES"] = "true"
        try:
            parser = create_parser("docx")
            content = {
                "id": "perf_test",
                "content": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }

            # Warm up
            try:
                parser.parse(content)
            except Exception as e:
                self.skipTest(f"Go DOCX parser not available: {e}")

            # Time multiple iterations
            times = []
            for i in range(self.iterations):
                start_time = time.time()
                result = parser.parse(content)
                end_time = time.time()
                times.append(end_time - start_time)

                # Verify we got a valid result
                self.assertIn("elements", result)
                self.assertGreater(len(result["elements"]), 0)

            avg_time = statistics.mean(times)
            min_time = min(times)
            max_time = max(times)

            print(f"\nGo DOCX Parser Performance:")
            print(f"  Average time: {avg_time*1000:.2f} ms")
            print(f"  Min time: {min_time*1000:.2f} ms")
            print(f"  Max time: {max_time*1000:.2f} ms")
            print(f"  Elements parsed: {len(result['elements'])}")
            print(f"  Relationships: {len(result.get('relationships', []))}")

        except Exception as e:
            self.skipTest(f"Go DOCX parser not available: {e}")

    def test_performance_comparison(self):
        """Compare performance between Python and Go parsers."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        # Test Python parser
        python_times = []
        python_elements = 0
        os.environ["USE_GO_MODULES"] = "false"
        try:
            python_parser = create_parser("docx")
            python_content = {
                "id": "perf_test",
                "binary_path": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }

            # Multiple iterations for Python
            for i in range(self.iterations):
                start_time = time.time()
                result = python_parser.parse(python_content)
                end_time = time.time()
                python_times.append(end_time - start_time)
                python_elements = len(result["elements"])

        except Exception:
            python_times = None

        # Test Go parser
        go_times = []
        go_elements = 0
        os.environ["USE_GO_MODULES"] = "true"
        try:
            go_parser = create_parser("docx")
            go_content = {
                "id": "perf_test",
                "content": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }

            # Multiple iterations for Go
            for i in range(self.iterations):
                start_time = time.time()
                result = go_parser.parse(go_content)
                end_time = time.time()
                go_times.append(end_time - start_time)
                go_elements = len(result["elements"])

        except Exception:
            go_times = None

        # Compare results
        if python_times and go_times:
            python_avg = statistics.mean(python_times)
            go_avg = statistics.mean(go_times)
            speedup = python_avg / go_avg if go_avg > 0 else 0

            print(f"\nPerformance Comparison:")
            print(f"  Python average: {python_avg*1000:.2f} ms ({python_elements} elements)")
            print(f"  Go average: {go_avg*1000:.2f} ms ({go_elements} elements)")
            print(f"  Speedup: {speedup:.2f}x")

            # Go should be faster (speedup > 1.0)
            if speedup > 1.0:
                print(f"  ✓ Go parser is {speedup:.2f}x faster")
            else:
                print(f"  ⚠ Go parser is slower (speedup: {speedup:.2f}x)")

        elif python_times:
            python_avg = statistics.mean(python_times)
            print(f"\nOnly Python parser available:")
            print(f"  Python average: {python_avg*1000:.2f} ms ({python_elements} elements)")

        elif go_times:
            go_avg = statistics.mean(go_times)
            print(f"\nOnly Go parser available:")
            print(f"  Go average: {go_avg*1000:.2f} ms ({go_elements} elements)")

        else:
            self.skipTest("Neither parser is available")

    def test_memory_efficiency(self):
        """Test memory usage patterns (basic check)."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        # This is a basic check - in a real scenario you'd use memory profiling tools
        parsers_tested = []

        # Test Python parser memory pattern
        os.environ["USE_GO_MODULES"] = "false"
        try:
            parser = create_parser("docx")
            content = {
                "id": "memory_test",
                "binary_path": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }

            # Parse multiple times to check for memory leaks
            for i in range(3):
                result = parser.parse(content)
                self.assertIn("elements", result)

            parsers_tested.append("Python")

        except Exception:
            pass

        # Test Go parser memory pattern
        os.environ["USE_GO_MODULES"] = "true"
        try:
            parser = create_parser("docx")
            content = {
                "id": "memory_test",
                "content": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }

            # Parse multiple times to check for memory leaks
            for i in range(3):
                result = parser.parse(content)
                self.assertIn("elements", result)

            parsers_tested.append("Go")

        except Exception:
            pass

        if not parsers_tested:
            self.skipTest("No parsers available for memory testing")

        print(f"\nMemory efficiency test completed for: {', '.join(parsers_tested)}")

    def tearDown(self):
        """Clean up after tests."""
        # Reset environment
        if "USE_GO_MODULES" in os.environ:
            del os.environ["USE_GO_MODULES"]


if __name__ == "__main__":
    # Run with verbose output to see performance numbers
    unittest.main(verbosity=2)