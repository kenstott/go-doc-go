"""
Unit tests for DOCX document parser.
"""

import os
import tempfile
import pytest
from typing import Dict, Any
import json

# Add parent directory to path
import sys
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__)))))

from go_doc_go.document_parser.docx import DocxParser
from go_doc_go.storage import ElementType
from go_doc_go.relationships import RelationshipType

# Check if python-docx is available
try:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    PYTHON_DOCX_AVAILABLE = True
except ImportError:
    PYTHON_DOCX_AVAILABLE = False


def create_simple_docx(path: str, title: str = "Test Document", 
                      paragraphs: list = None, add_table: bool = False,
                      add_list: bool = False, add_image: bool = False) -> None:
    """Create a simple DOCX file for testing."""
    if not PYTHON_DOCX_AVAILABLE:
        pytest.skip("python-docx not available")
    
    doc = Document()
    
    # Add title
    doc.add_heading(title, 0)
    
    # Add paragraphs
    if paragraphs:
        for para_text in paragraphs:
            p = doc.add_paragraph(para_text)
    
    # Add list if requested
    if add_list:
        doc.add_paragraph("Here is a list:", style='Body Text')
        doc.add_paragraph('First item', style='List Bullet')
        doc.add_paragraph('Second item', style='List Bullet')
        doc.add_paragraph('Third item', style='List Bullet')
    
    # Add table if requested
    if add_table:
        doc.add_paragraph("Table example:", style='Body Text')
        table = doc.add_table(rows=3, cols=3)
        table.style = 'Table Grid'
        
        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Header 1'
        hdr_cells[1].text = 'Header 2'
        hdr_cells[2].text = 'Header 3'
        
        # Data rows
        for i in range(1, 3):
            row_cells = table.rows[i].cells
            for j in range(3):
                row_cells[j].text = f'Row {i}, Col {j+1}'
    
    doc.save(path)


def create_styled_docx(path: str) -> None:
    """Create a DOCX with various styles and formatting."""
    if not PYTHON_DOCX_AVAILABLE:
        pytest.skip("python-docx not available")
    
    doc = Document()
    
    # Title
    doc.add_heading('Styled Document', 0)
    
    # Various heading levels
    doc.add_heading('Heading Level 1', 1)
    doc.add_paragraph('Content under heading 1')
    
    doc.add_heading('Heading Level 2', 2)
    doc.add_paragraph('Content under heading 2')
    
    # Formatted text
    p = doc.add_paragraph()
    p.add_run('Normal text. ')
    p.add_run('Bold text. ').bold = True
    p.add_run('Italic text. ').italic = True
    
    # Add hyperlink-style text (not actual hyperlink but formatted like one)
    p = doc.add_paragraph()
    run = p.add_run('This looks like a link')
    run.font.color.rgb = RGBColor(0, 0, 255)
    run.font.underline = True
    
    # Code block (using monospace font)
    p = doc.add_paragraph()
    run = p.add_run('def hello_world():\n    print("Hello, World!")')
    run.font.name = 'Courier New'
    run.font.size = Pt(10)
    
    # Numbered list
    doc.add_paragraph('First numbered item', style='List Number')
    doc.add_paragraph('Second numbered item', style='List Number')
    
    doc.save(path)


def create_multipage_docx(path: str, num_sections: int = 3) -> None:
    """Create a multi-section DOCX document."""
    if not PYTHON_DOCX_AVAILABLE:
        pytest.skip("python-docx not available")
    
    doc = Document()
    
    for section_num in range(num_sections):
        # Add section heading
        doc.add_heading(f'Section {section_num + 1}', 1)
        
        # Add some paragraphs
        for para_num in range(3):
            doc.add_paragraph(
                f'This is paragraph {para_num + 1} in section {section_num + 1}. '
                f'It contains some sample text to demonstrate parsing capabilities.'
            )
        
        # Add page break between sections (except last)
        if section_num < num_sections - 1:
            doc.add_page_break()
    
    doc.save(path)


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not available")
class TestDocxParser:
    """Test suite for DOCX parser."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.parser = DocxParser()
    
    def test_parser_initialization(self):
        """Test DOCX parser initialization."""
        # Default initialization
        parser1 = DocxParser()
        assert hasattr(parser1, 'extract_styles')
        assert hasattr(parser1, 'extract_comments')
        
        # Custom configuration
        config = {
            "extract_styles": False,
            "extract_comments": False,
            "extract_headers_footers": False,
            "max_content_preview": 50
        }
        parser2 = DocxParser(config)
        assert parser2.extract_styles == False
        assert parser2.extract_comments == False
    
    def test_basic_docx_parsing(self):
        """Test basic DOCX parsing functionality."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create a simple DOCX
            create_simple_docx(tmp_path, 
                             title="Test DOCX Document",
                             paragraphs=["This is the first paragraph.",
                                       "This is the second paragraph with more content."])
            
            # Parse it
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {
                    "doc_id": "test_docx_123"
                }
            }
            
            result = self.parser.parse(content)
            
            # Check basic structure
            assert "document" in result
            assert "elements" in result
            assert "relationships" in result
            
            # Check document
            doc = result["document"]
            assert doc["doc_id"] == "test_docx_123"
            assert doc["doc_type"] == "docx"
            
            # Check elements
            elements = result["elements"]
            assert len(elements) > 0
            
            # Check for root element
            root = next((e for e in elements if e["element_type"] == ElementType.ROOT.value), None)
            assert root is not None
            
            # Check for paragraphs
            paragraphs = [e for e in elements if e["element_type"] == ElementType.PARAGRAPH.value]
            assert len(paragraphs) >= 2  # At least the two paragraphs we added
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_text_extraction(self):
        """Test text extraction from DOCX."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create DOCX with specific text
            create_simple_docx(tmp_path,
                             title="Document Title",
                             paragraphs=[
                                 "This is a paragraph with some text content.",
                                 "It has multiple paragraphs.",
                                 "Another paragraph here with different content."
                             ])
            
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = self.parser.parse(content)
            elements = result["elements"]
            
            # Should have extracted paragraphs
            paragraphs = [e for e in elements if e["element_type"] == ElementType.PARAGRAPH.value]
            assert len(paragraphs) > 0
            
            # Check content preservation
            content_texts = [p["content_preview"] for p in paragraphs]
            all_content = " ".join(content_texts)
            
            assert "paragraph with some text" in all_content
            assert "multiple paragraphs" in all_content
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_header_extraction(self):
        """Test extraction of headers/headings."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create DOCX with headers
            doc = Document()
            doc.add_heading('Main Title', 0)
            doc.add_heading('Chapter 1', 1)
            doc.add_paragraph('Chapter 1 content')
            doc.add_heading('Section 1.1', 2)
            doc.add_paragraph('Section content')
            doc.save(tmp_path)
            
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = self.parser.parse(content)
            elements = result["elements"]
            
            # Should have headers
            headers = [e for e in elements if e["element_type"] == ElementType.HEADER.value]
            assert len(headers) >= 3  # Main title, Chapter 1, Section 1.1
            
            # Check header levels
            header_texts = [h["content_preview"] for h in headers]
            assert any("Main Title" in text for text in header_texts)
            assert any("Chapter 1" in text for text in header_texts)
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_list_extraction(self):
        """Test extraction of lists."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create DOCX with lists
            create_simple_docx(tmp_path,
                             title="Document with Lists",
                             paragraphs=["Some text before the list."],
                             add_list=True)
            
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = self.parser.parse(content)
            elements = result["elements"]
            
            # Should have list items
            list_items = [e for e in elements if e["element_type"] == ElementType.LIST_ITEM.value]
            # Lists might be parsed as paragraphs depending on implementation
            paragraphs = [e for e in elements if e["element_type"] == ElementType.PARAGRAPH.value]
            
            # Check that list content exists somewhere
            all_elements = list_items + paragraphs
            content_texts = [e["content_preview"] for e in all_elements]
            all_content = " ".join(content_texts)
            
            assert "First item" in all_content
            assert "Second item" in all_content
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_table_extraction(self):
        """Test extraction of tables."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create DOCX with table
            create_simple_docx(tmp_path,
                             title="Document with Table",
                             paragraphs=["Some text before the table."],
                             add_table=True)
            
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = self.parser.parse(content)
            elements = result["elements"]
            
            # Should have table elements
            tables = [e for e in elements if e["element_type"] == ElementType.TABLE.value]
            table_cells = [e for e in elements if e["element_type"] == ElementType.TABLE_CELL.value]
            
            # Check that table content exists
            assert len(tables) > 0 or len(table_cells) > 0
            
            # Check cell content
            if table_cells:
                cell_texts = [c["content_preview"] for c in table_cells]
                all_cells = " ".join(cell_texts)
                assert "Header" in all_cells or "Row" in all_cells
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_styled_text_extraction(self):
        """Test extraction of styled text (bold, italic, etc)."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create styled DOCX
            create_styled_docx(tmp_path)
            
            parser = DocxParser({"extract_styles": True})
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = parser.parse(content)
            elements = result["elements"]
            
            # Should have various elements
            paragraphs = [e for e in elements if e["element_type"] == ElementType.PARAGRAPH.value]
            assert len(paragraphs) > 0
            
            # Check that styled content is preserved
            content_texts = [p["content_preview"] for p in paragraphs]
            all_content = " ".join(content_texts)
            
            assert "Bold text" in all_content
            assert "Italic text" in all_content
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_multipage_document(self):
        """Test parsing of multi-section documents."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create multi-section DOCX
            create_multipage_docx(tmp_path, num_sections=3)
            
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = self.parser.parse(content)
            elements = result["elements"]
            
            # Should have multiple sections worth of content
            headers = [e for e in elements if e["element_type"] == ElementType.HEADER.value]
            paragraphs = [e for e in elements if e["element_type"] == ElementType.PARAGRAPH.value]
            
            assert len(headers) >= 3  # One per section
            assert len(paragraphs) >= 9  # 3 per section
            
            # Check section headers
            header_texts = [h["content_preview"] for h in headers]
            assert any("Section 1" in text for text in header_texts)
            assert any("Section 2" in text for text in header_texts)
            assert any("Section 3" in text for text in header_texts)
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_empty_docx(self):
        """Test handling of empty DOCX."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
        
        try:
            # Create empty DOCX
            doc = Document()
            doc.save(tmp_path)
            
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            result = self.parser.parse(content)
            
            # Should handle empty DOCX
            assert result is not None
            assert "document" in result
            assert "elements" in result
            
            # Should have at least root element
            elements = result["elements"]
            root = next((e for e in elements if e["element_type"] == ElementType.ROOT.value), None)
            assert root is not None
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
    
    def test_binary_content_handling(self):
        """Test handling of binary content without file path."""
        # Create a simple DOCX in memory
        doc = Document()
        doc.add_heading("Test Document", 0)
        doc.add_paragraph("Test content")
        
        # Save to bytes
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            tmp_path = tmp.name
            doc.save(tmp_path)
            
            with open(tmp_path, 'rb') as f:
                docx_bytes = f.read()
            
            os.unlink(tmp_path)
        
        content = {
            "id": "memory_docx",
            "content": docx_bytes,
            "metadata": {}
        }
        
        result = self.parser.parse(content)
        
        # Should handle binary content
        assert result is not None
        assert "document" in result
        
        # Should have created elements
        elements = result["elements"]
        paragraphs = [e for e in elements if e["element_type"] == ElementType.PARAGRAPH.value]
        assert len(paragraphs) > 0


@pytest.mark.skipif(not PYTHON_DOCX_AVAILABLE, reason="python-docx not available")
class TestDocxParserErrorHandling:
    """Test error handling in DOCX parser."""
    
    def setup_method(self):
        """Set up test fixtures."""
        self.parser = DocxParser()
    
    def test_nonexistent_file(self):
        """Test handling of non-existent file."""
        content = {
            "id": "/nonexistent/file.docx",
            "binary_path": "/nonexistent/file.docx",
            "metadata": {}
        }
        
        with pytest.raises(Exception):
            self.parser.parse(content)
    
    def test_corrupt_docx(self):
        """Test handling of corrupt DOCX file."""
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
            # Write invalid DOCX content
            tmp.write(b"This is not a valid DOCX file")
            tmp_path = tmp.name
        
        try:
            content = {
                "id": tmp_path,
                "binary_path": tmp_path,
                "metadata": {}
            }
            
            with pytest.raises(Exception):
                self.parser.parse(content)
            
        finally:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)


if __name__ == "__main__":
    pytest.main([__file__, "-v"])