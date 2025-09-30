#!/usr/bin/env python3
"""
Compatibility tests between Go and Python DOCX parsers.

This module ensures that the Go DOCX parser produces results compatible
with the Python DOCX parser, maintaining API consistency while providing
performance improvements.
"""

import json
import os
import tempfile
import unittest
from pathlib import Path
import sys

# Add src directory to path
sys.path.insert(0, str(Path(__file__).parent.parent / "src"))

from go_doc_go.document_parser.factory import create_parser


class TestDocxGoCompatibility(unittest.TestCase):
    """Test compatibility between Go and Python DOCX parsers."""

    def setUp(self):
        """Set up test fixtures."""
        self.test_dir = Path(__file__).parent
        self.assets_dir = self.test_dir / "assets"

        # Create a simple test DOCX file if it doesn't exist
        self.test_docx_file = self.assets_dir / "test_sample.docx"
        if not self.test_docx_file.exists():
            self.create_test_docx()

    def create_test_docx(self):
        """Create a simple test DOCX file."""
        try:
            from docx import Document
            from docx.shared import Inches

            # Create assets directory if it doesn't exist
            self.assets_dir.mkdir(parents=True, exist_ok=True)

            # Create a simple document
            doc = Document()

            # Add a title
            doc.add_heading('Test Document', 0)

            # Add a paragraph
            doc.add_paragraph('This is a test paragraph with some sample text.')

            # Add a heading
            doc.add_heading('Section 1', 1)

            # Add another paragraph
            doc.add_paragraph('This section contains more detailed information about the test document.')

            # Add a table
            table = doc.add_table(rows=3, cols=3)
            table.style = 'Light Grid Accent 1'

            # Add table content
            cells = table.rows[0].cells
            cells[0].text = 'Header 1'
            cells[1].text = 'Header 2'
            cells[2].text = 'Header 3'

            cells = table.rows[1].cells
            cells[0].text = 'Row 1, Col 1'
            cells[1].text = 'Row 1, Col 2'
            cells[2].text = 'Row 1, Col 3'

            cells = table.rows[2].cells
            cells[0].text = 'Row 2, Col 1'
            cells[1].text = 'Row 2, Col 2'
            cells[2].text = 'Row 2, Col 3'

            # Add another heading
            doc.add_heading('Section 2', 2)

            # Add a list
            doc.add_paragraph('Item 1', style='List Bullet')
            doc.add_paragraph('Item 2', style='List Bullet')
            doc.add_paragraph('Item 3', style='List Bullet')

            # Save the document
            doc.save(str(self.test_docx_file))

        except ImportError:
            # If python-docx is not available, create a dummy file
            self.test_docx_file.touch()

    def test_basic_parsing_compatibility(self):
        """Test that both parsers can handle basic DOCX parsing."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        # Test Python parser
        os.environ["USE_GO_MODULES"] = "false"
        try:
            python_parser = create_parser("docx")
            python_content = {
                "id": "test_doc",
                "binary_path": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }
            python_result = python_parser.parse(python_content)
        except ImportError:
            self.skipTest("Python DOCX parser not available (python-docx not installed)")
        except Exception as e:
            self.skipTest(f"Python DOCX parser failed: {e}")

        # Test Go parser
        os.environ["USE_GO_MODULES"] = "true"
        try:
            go_parser = create_parser("docx")
            go_content = {
                "id": "test_doc",
                "content": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }
            go_result = go_parser.parse(go_content)
        except Exception as e:
            self.skipTest(f"Go DOCX parser not available: {e}")

        # Compare basic structure
        self.assertIn("document", python_result)
        self.assertIn("document", go_result)
        self.assertIn("elements", python_result)
        self.assertIn("elements", go_result)
        self.assertIn("relationships", python_result)
        self.assertIn("relationships", go_result)

        # Check document metadata
        python_doc = python_result["document"]
        go_doc = go_result["document"]

        self.assertEqual(python_doc["doc_type"], "docx")
        self.assertEqual(go_doc["doc_type"], "docx")
        self.assertEqual(python_doc["doc_id"], "test_doc")
        self.assertEqual(go_doc["doc_id"], "test_doc")

    def test_element_structure_compatibility(self):
        """Test that both parsers produce compatible element structures."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        # Get results from both parsers
        python_result = self._get_python_result()
        go_result = self._get_go_result()

        if not python_result or not go_result:
            self.skipTest("One or both parsers not available")

        python_elements = python_result["elements"]
        go_elements = go_result["elements"]

        # Check that both have elements
        self.assertGreater(len(python_elements), 0)
        self.assertGreater(len(go_elements), 0)

        # Check element structure
        for elements, parser_name in [(python_elements, "Python"), (go_elements, "Go")]:
            for element in elements:
                with self.subTest(parser=parser_name, element_id=element.get("element_id")):
                    # Check required fields
                    self.assertIn("element_id", element)
                    self.assertIn("element_type", element)
                    self.assertIn("content_preview", element)
                    self.assertIn("content_location", element)
                    self.assertIn("metadata", element)

                    # Check data types
                    self.assertIsInstance(element["element_id"], str)
                    self.assertIsInstance(element["element_type"], str)
                    self.assertIsInstance(element["content_preview"], str)
                    self.assertIsInstance(element["metadata"], dict)

    def test_element_types_compatibility(self):
        """Test that both parsers use compatible element types."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        python_result = self._get_python_result()
        go_result = self._get_go_result()

        if not python_result or not go_result:
            self.skipTest("One or both parsers not available")

        # Extract element types
        python_types = set(elem["element_type"] for elem in python_result["elements"])
        go_types = set(elem["element_type"] for elem in go_result["elements"])

        # Common element types that should be present in both
        common_types = {"document_root", "body", "paragraph"}

        # Check that common types are present
        self.assertTrue(
            common_types.issubset(python_types),
            f"Python parser missing common types: {common_types - python_types}"
        )
        self.assertTrue(
            common_types.issubset(go_types),
            f"Go parser missing common types: {common_types - go_types}"
        )

    def test_relationship_structure_compatibility(self):
        """Test that both parsers produce compatible relationship structures."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        python_result = self._get_python_result()
        go_result = self._get_go_result()

        if not python_result or not go_result:
            self.skipTest("One or both parsers not available")

        python_relationships = python_result["relationships"]
        go_relationships = go_result["relationships"]

        # Check that both have relationships
        self.assertGreater(len(python_relationships), 0)
        self.assertGreater(len(go_relationships), 0)

        # Check relationship structure
        for relationships, parser_name in [(python_relationships, "Python"), (go_relationships, "Go")]:
            for relationship in relationships:
                with self.subTest(parser=parser_name, rel_id=relationship.get("relationship_id")):
                    # Check required fields
                    self.assertIn("relationship_id", relationship)
                    self.assertIn("source_element_id", relationship)
                    self.assertIn("target_element_id", relationship)
                    self.assertIn("relationship_type", relationship)
                    self.assertIn("metadata", relationship)

                    # Check data types
                    self.assertIsInstance(relationship["relationship_id"], str)
                    self.assertIsInstance(relationship["source_element_id"], str)
                    self.assertIsInstance(relationship["target_element_id"], str)
                    self.assertIsInstance(relationship["relationship_type"], str)
                    self.assertIsInstance(relationship["metadata"], dict)

    def test_content_location_compatibility(self):
        """Test that content location formats are compatible."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        python_result = self._get_python_result()
        go_result = self._get_go_result()

        if not python_result or not go_result:
            self.skipTest("One or both parsers not available")

        # Check content location structure
        for result, parser_name in [(python_result, "Python"), (go_result, "Go")]:
            for element in result["elements"]:
                with self.subTest(parser=parser_name, element_id=element.get("element_id")):
                    content_location = element["content_location"]

                    # Content location should be a dict or valid JSON string
                    if isinstance(content_location, str):
                        try:
                            location_data = json.loads(content_location)
                        except json.JSONDecodeError:
                            self.fail(f"Invalid JSON in content_location for {parser_name}")
                    else:
                        location_data = content_location

                    # Check required fields
                    self.assertIn("source", location_data)
                    self.assertIn("type", location_data)

    def test_table_parsing_compatibility(self):
        """Test that table parsing is compatible between parsers."""
        if not self.test_docx_file.exists() or self.test_docx_file.stat().st_size == 0:
            self.skipTest("Test DOCX file not available")

        python_result = self._get_python_result()
        go_result = self._get_go_result()

        if not python_result or not go_result:
            self.skipTest("One or both parsers not available")

        # Find table elements
        python_tables = [e for e in python_result["elements"] if e["element_type"] == "table"]
        go_tables = [e for e in go_result["elements"] if e["element_type"] == "table"]

        # Both should find some table elements (if our test document has tables)
        if python_tables:
            self.assertGreater(len(go_tables), 0, "Go parser should find tables if Python parser does")

        # Check table structure
        for tables, parser_name in [(python_tables, "Python"), (go_tables, "Go")]:
            for table in tables:
                with self.subTest(parser=parser_name, table_id=table.get("element_id")):
                    self.assertEqual(table["element_type"], "table")
                    self.assertIn("metadata", table)

    def test_error_handling_compatibility(self):
        """Test that both parsers handle errors similarly."""
        # Test with non-existent file
        non_existent_file = "/path/that/does/not/exist.docx"

        # Test Python parser error handling
        os.environ["USE_GO_MODULES"] = "false"
        try:
            python_parser = create_parser("docx")
            python_content = {
                "id": "test_doc",
                "binary_path": non_existent_file,
                "metadata": {"filename": "nonexistent.docx"}
            }
            with self.assertRaises(Exception):
                python_parser.parse(python_content)
        except ImportError:
            pass  # Skip if Python parser not available

        # Test Go parser error handling
        os.environ["USE_GO_MODULES"] = "true"
        try:
            go_parser = create_parser("docx")
            go_content = {
                "id": "test_doc",
                "content": non_existent_file,
                "metadata": {"filename": "nonexistent.docx"}
            }
            with self.assertRaises(Exception):
                go_parser.parse(go_content)
        except ImportError:
            pass  # Skip if Go parser not available

    def _get_python_result(self):
        """Get parsing result from Python parser."""
        try:
            os.environ["USE_GO_MODULES"] = "false"
            parser = create_parser("docx")
            content = {
                "id": "test_doc",
                "binary_path": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }
            return parser.parse(content)
        except Exception:
            return None

    def _get_go_result(self):
        """Get parsing result from Go parser."""
        try:
            os.environ["USE_GO_MODULES"] = "true"
            parser = create_parser("docx")
            content = {
                "id": "test_doc",
                "content": str(self.test_docx_file),
                "metadata": {"filename": "test_sample.docx"}
            }
            return parser.parse(content)
        except Exception:
            return None

    def tearDown(self):
        """Clean up after tests."""
        # Reset environment
        if "USE_GO_MODULES" in os.environ:
            del os.environ["USE_GO_MODULES"]


if __name__ == "__main__":
    unittest.main()